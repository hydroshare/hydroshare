"""
Generate the HydroShare Content Security Audit executive summary as a Word document.
Usage: python3 scripts/generate_report_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Styles ────────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # medium blue
    return p

def add_para(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_table_row(table, cells, bold=False, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=bold)
        if shade:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shade)
            tcPr.append(shd)
    return row


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run('HydroShare Content Security Audit')
r.font.name = 'Calibri'
r.font.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(16)
r2 = sub.add_run('Executive Summary  ·  July 2026')
r2.font.name = 'Calibri'
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
# METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Methodology', level=2)
add_para(doc,
    'This audit was conducted using an AI-assisted review tool (GPT-4o-mini via GitHub Models API) '
    'developed and operated by CUAHSI staff. The tool scanned resource titles, abstracts, and metadata '
    'and returned classifications that were reviewed, validated, and refined through iterative human '
    'oversight before finalizing findings. All categorization decisions and threshold adjustments were '
    'made under human direction.')

# ══════════════════════════════════════════════════════════════════════════════
# SCOPE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Scope', level=2)
add_para(doc,
    '20,616 resources were scanned across all visibility tiers — 8,136 public, 909 discoverable-only, '
    'and 11,550 private.')

# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC RISK
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Are Public Resources at Risk?', level=2)
add_para(doc,
    'No. The 19 malicious public resources identified have been removed. These were commercial SEO spam — '
    'Indonesian travel and car rental ads, PUBG gaming account listings, and health/wellness marketing — '
    'posted publicly to gain search engine backlinks, not to target HydroShare users. The '
    'discoverable-only catalog was entirely clean.')

# ══════════════════════════════════════════════════════════════════════════════
# SECURITY FINDING
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Security Finding: Confirmed Malicious Content', level=2)
add_para(doc,
    'The scan identified 253 malicious resources created by 189 accounts with no connection to water '
    'science. These accounts registered on HydroShare and used it as a free hosting platform for '
    'commercial spam. Content types found:')

spam_types = [
    'Pharmaceutical ads (erectile dysfunction medications, controlled substances)',
    'Escort and adult services',
    'Online gambling promotions',
    'Travel booking and visa services',
    'Tech support and financial services scams',
    'General commercial advertising (construction, retail, education)',
]
for s in spam_types:
    add_bullet(doc, s)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Recommended action: ')
set_font(r, bold=True)
r2 = p.add_run(
    'The 189 accounts and their 253 resources are listed in spam_security_report.csv. '
    'Suspend all 189 accounts.')
set_font(r2)

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT QUALITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Content Quality Findings (Not Security Concerns)', level=2)
add_para(doc,
    'The scan also flagged approximately 2,500 resources across six quality categories. These represent '
    'normal platform activity and require no security action.')

# Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.autofit = True

hdr = table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Count'
hdr[2].text = 'Description'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D6E4F0')
    tcPr.append(shd)

rows_data = [
    ('Ambiguous / uncategorized content', '963',
     'Orphaned workshop data, personal uploads, low-quality titles'),
    ('Developer and researcher test resources', '654',
     'JupyterHub demo notebooks, temp files, prototype resources'),
    ('Geographic placeholder entries', '475',
     'Single watershed or location names with no scientific content'),
    ('Student homework and class exercises', '157',
     'CEE 6400, Snow Modeling HW, Cyber Carpentry assignments'),
    ('Off-topic academic content', '114',
     'Non-water research (geology, ecology, finance) uploaded by researchers'),
    ('Platform infrastructure assets', '111',
     'Tool logos, JupyterHub screenshots, release notes'),
]
for i, (cat, cnt, desc) in enumerate(rows_data):
    row = table.add_row()
    row.cells[0].text = cat
    row.cells[1].text = cnt
    row.cells[2].text = desc
    shade = 'F2F8FD' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade)
        tcPr.append(shd)

doc.add_paragraph()
add_para(doc,
    'These are housekeeping items. If a content quality cleanup is desired, it can be addressed '
    'separately as a platform maintenance exercise.')

# ══════════════════════════════════════════════════════════════════════════════
# BOTTOM LINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Bottom Line', level=2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
for text, bold in [
    ('Public content is clean. ', True),
    ('253 malicious resources from 189 bot/spam accounts were found, primarily in private storage. '
     'Suspending the 189 flagged accounts closes the security issue. The remaining flagged content '
     'is normal academic platform clutter, not abuse.', False),
]:
    r = p.add_run(text)
    set_font(r, bold=bold)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(f'HydroShare Security Audit · CUAHSI · July 2026 · Confidential')
fr.font.name = 'Calibri'
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

# ── Save ──────────────────────────────────────────────────────────────────────
output = 'HydroShare_Security_Audit_Executive_Summary.docx'
doc.save(output)
print(f'Saved: {output}')
